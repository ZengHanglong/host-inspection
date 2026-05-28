"""
Report Generator - 生成主机巡检报告 DOCX
"""
from datetime import datetime
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ALERT_LABELS = {
    "host_status": "主机状态异常",
    "snapshot_expired": "过期快照",
    "large_vm": "大容量虚拟机",
    "naming_rule": "命名/备注不规范",
    "idle_vm": "闲置资产",
}


PERIODIC_ITEM_DETAILS = {
    "snapshot": {
        "result_title": "过期快照清单",
        "headers": ["虚拟机", "集群", "快照", "天数", "级别", "建议动作"],
        "build_rows": lambda platform: [
            [
                row.get("vm_name", "--"),
                row.get("cluster_name", "--"),
                row.get("snapshot_name", "root-snapshot"),
                str(row.get("snapshot_days", 0)),
                row.get("status", "warning"),
                row.get("action", "--"),
            ]
            for row in platform.get("expired_snapshots", [])
        ],
    },
    "naming": {
        "result_title": "命名与备注问题清单",
        "headers": ["虚拟机", "集群", "IP", "问题", "当前备注"],
        "build_rows": lambda platform: [
            [
                row.get("vm_name", "--"),
                row.get("cluster_name", "--"),
                row.get("ip", "--"),
                "；".join(row.get("issues", [])) or "--",
                row.get("current_note", "") or "--",
            ]
            for row in platform.get("naming_issues", [])
        ],
    },
    "idle_vm": {
        "result_title": "闲置资产清单",
        "headers": ["虚拟机", "集群", "电源状态", "创建天数", "容量(GB)", "建议动作"],
        "build_rows": lambda platform: [
            [
                row.get("vm_name", "--"),
                row.get("cluster_name", "--"),
                row.get("power_state", "--"),
                str(row.get("created_days", 0)),
                str(row.get("disk_gb", 0)),
                row.get("action", "--"),
            ]
            for row in platform.get("idle_vms", [])
        ],
    },
    "large_vm": {
        "result_title": "大容量虚拟机清单",
        "headers": ["虚拟机", "集群", "容量(TB)", "CPU", "内存(GB)", "建议动作"],
        "build_rows": lambda platform: [
            [
                row.get("vm_name", "--"),
                row.get("cluster_name", "--"),
                str(row.get("disk_tb", 0)),
                str(row.get("cpu", 0)),
                str(row.get("memory_gb", 0)),
                row.get("action", "--"),
            ]
            for row in platform.get("large_vms", [])
        ],
    },
}


def create_inspection_report(
    company_name: str = "厦门国际信托",
    report_date: Optional[str] = None,
    data: Optional[Dict[str, Any]] = None,
    thresholds: Optional[Dict[str, Any]] = None,
    generator: str = "厦门汉为软件技术有限公司",
) -> Document:
    if report_date is None:
        report_date = datetime.now().strftime("%Y%m%d")

    report = data or {}
    platforms = report.get("platforms", {})
    doc = Document()

    _set_base_style(doc)
    _add_cover(doc, company_name, report_date, generator, report)
    _add_executive_summary(doc, report)
    _add_scope_section(doc, report, thresholds or {})

    if platforms:
        doc.add_heading("三、平台巡检结果", level=1)
        for platform_code, platform in platforms.items():
            _add_platform_section(doc, platform_code, platform)
    else:
        doc.add_heading("三、平台巡检结果", level=1)
        doc.add_paragraph("当前无已接通平台，未生成任何真实巡检结果。")

    _add_alert_section(doc, report.get("warnings", []))
    _add_error_section(doc, report.get("errors", []))
    _add_data_source_section(doc, report)
    return doc


def _set_base_style(doc: Document) -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(10.5)

    for style_name, font_size in (("Title", 20), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(font_size)
        style.font.bold = True


def _add_cover(
    doc: Document,
    company_name: str,
    report_date: str,
    generator: str,
    report: Dict[str, Any],
) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, company_name, size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(subtitle, "主机巡检报告", size=18, bold=True)

    report_type = report.get("report_type") or "每日巡检报告"
    type_paragraph = doc.add_paragraph()
    type_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(type_paragraph, report_type, size=12)

    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(date_paragraph, f"报告日期：{_format_report_date(report_date)}", size=11)

    generator_paragraph = doc.add_paragraph()
    generator_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(generator_paragraph, f"生成单位：{generator}", size=11)

    declaration = doc.add_paragraph()
    declaration.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        declaration,
        report.get("data_declaration") or "本报告所有数据均来自真实API调用，无模拟数据。",
        size=10.5,
    )

    doc.add_page_break()


def _add_executive_summary(doc: Document, report: Dict[str, Any]) -> None:
    doc.add_heading("一、执行摘要", level=1)
    summary = report.get("overall_summary", {})
    statistics = report.get("overall_statistics", {})

    doc.add_paragraph(
        "本次巡检覆盖已接入平台的集群、主机和虚拟机运行状态，并同步输出过期快照、命名与备注维护、闲置资产、大容量虚拟机四类周期巡检结果。"
    )
    doc.add_paragraph(
        f"本次共纳管 {summary.get('total_platforms', 0)} 个平台，已配置 {summary.get('configured_platforms', 0)} 个，已连接 {summary.get('connected_platforms', 0)} 个。"
    )
    doc.add_paragraph(
        f"主机总数 {statistics.get('total_hosts', 0)} 台，其中正常 {statistics.get('normal_hosts', 0)} 台，警告 {statistics.get('warning_hosts', 0)} 台，严重 {statistics.get('critical_hosts', 0)} 台；当前共识别告警 {statistics.get('total_alerts', 0)} 条。"
    )

    key_findings = _collect_key_findings(report)
    if key_findings:
        doc.add_paragraph("本次重点发现：")
        for finding in key_findings:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(finding)
    else:
        doc.add_paragraph("本次未发现需要升级处理的重点异常。")


def _add_scope_section(doc: Document, report: Dict[str, Any], thresholds: Dict[str, Any]) -> None:
    doc.add_heading("二、巡检范围与判定口径", level=1)
    doc.add_paragraph("日巡检重点关注集群状态、主机资源使用率与平台连接状态；周期巡检重点关注快照、命名与备注、闲置资产和大容量虚拟机。")

    summary = report.get("overall_summary", {})
    doc.add_paragraph(
        f"平台覆盖范围：总平台 {summary.get('total_platforms', 0)} 个，当前生成报告的平台以已连接平台为准。"
    )

    cpu_warning = thresholds.get("cpu_warning", 70.0)
    cpu_critical = thresholds.get("cpu_critical", 80.0)
    memory_warning = thresholds.get("memory_warning", 70.0)
    memory_critical = thresholds.get("memory_critical", 80.0)
    storage_warning = thresholds.get("storage_warning", 60.0)
    storage_critical = thresholds.get("storage_critical", 70.0)

    threshold_table = _create_table(doc, ["指标", "警告阈值", "严重阈值"], [
        ["CPU使用率", f"{cpu_warning}%", f"{cpu_critical}%"],
        ["内存使用率", f"{memory_warning}%", f"{memory_critical}%"],
        ["存储使用率", f"{storage_warning}%", f"{storage_critical}%"],
    ])
    threshold_table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _add_platform_section(doc: Document, platform_code: str, platform: Dict[str, Any]) -> None:
    doc.add_heading(platform.get('name', platform_code), level=2)

    doc.add_paragraph(
        f"平台状态：{platform.get('status', '未知')}；数据来源：{platform.get('data_source', '无')}。"
    )
    _add_platform_summary_table(doc, platform)
    _add_cluster_section(doc, platform)
    _add_host_section(doc, platform)
    _add_periodic_section(doc, platform)


def _add_platform_summary_table(doc: Document, platform: Dict[str, Any]) -> None:
    stats = platform.get("statistics", {})
    _create_table(
        doc,
        ["指标", "值"],
        [
            ["集群数", str(stats.get("clusters", 0))],
            ["主机数", str(stats.get("total", 0))],
            ["虚拟机数", str(stats.get("vms", 0))],
            ["正常主机", str(stats.get("normal", 0))],
            ["警告主机", str(stats.get("warning", 0))],
            ["严重主机", str(stats.get("critical", 0))],
        ],
    )


def _add_cluster_section(doc: Document, platform: Dict[str, Any]) -> None:
    doc.add_heading("3.1 集群状态", level=2)
    clusters = platform.get("clusters", [])
    if not clusters:
        doc.add_paragraph("当前平台未返回集群数据。")
        return

    rows = []
    for cluster in clusters:
        rows.append([
            cluster.get("cluster_name", "--"),
            cluster.get("status", "unknown"),
            str(cluster.get("host_count", 0)),
            _format_percent(cluster.get("cpu_usage_percent", 0)),
            _format_percent(cluster.get("memory_usage_percent", 0)),
        ])

    _create_table(doc, ["集群", "状态", "主机数", "CPU", "内存"], rows)


def _add_host_section(doc: Document, platform: Dict[str, Any]) -> None:
    doc.add_heading("3.2 主机运行情况", level=2)
    hosts = platform.get("hosts", [])
    if not hosts:
        doc.add_paragraph("当前平台未返回主机数据。")
        return

    rows = []
    for host in hosts:
        rows.append([
            host.get("host_name", "--"),
            host.get("cluster_name", "--"),
            host.get("ip_address", "--"),
            _format_percent(host.get("cpu_usage_percent", 0)),
            _format_percent(host.get("memory_usage_percent", 0)),
            host.get("status", "unknown"),
        ])

    _create_table(doc, ["主机", "集群", "IP", "CPU", "内存", "状态"], rows)


def _add_periodic_section(doc: Document, platform: Dict[str, Any]) -> None:
    doc.add_heading("3.3 周期巡检项", level=2)
    periodic_items = platform.get("periodic_items", [])
    if not periodic_items:
        doc.add_paragraph("当前平台暂无周期巡检项结果。")
        return

    for item in periodic_items:
        doc.add_heading(item.get("name", "周期巡检项"), level=3)
        doc.add_paragraph(f"巡检目标：{item.get('target', '--')}")
        doc.add_paragraph(f"检查重点：{item.get('focus', '--')}")
        doc.add_paragraph(f"结果结论：本次共发现 {item.get('count', 0)} 项问题。")
        doc.add_paragraph(_build_periodic_recommendation(item))

        detail = PERIODIC_ITEM_DETAILS.get(item.get("code"))
        if not detail:
            continue

        rows = detail["build_rows"](platform)
        if rows:
            doc.add_paragraph(detail["result_title"])
            _create_table(doc, detail["headers"], rows)
        else:
            doc.add_paragraph("本次未发现相关异常。")


def _add_alert_section(doc: Document, warnings: List[Dict[str, Any]]) -> None:
    doc.add_heading("四、告警汇总", level=1)
    if not warnings:
        doc.add_paragraph("本次未产生告警。")
        return

    rows = []
    for alert in warnings:
        rows.append([
            alert.get("platform_name", "--"),
            alert.get("host_name", "--"),
            ALERT_LABELS.get(alert.get("alert_type"), alert.get("alert_type", "--")),
            alert.get("alert_level", "warning"),
            alert.get("message", "--"),
        ])

    _create_table(doc, ["平台", "对象", "类型", "级别", "详情"], rows)


def _add_error_section(doc: Document, errors: List[Dict[str, Any]]) -> None:
    doc.add_heading("五、未接通平台说明", level=1)
    if not errors:
        doc.add_paragraph("当前无连接失败的平台。")
        return

    rows = [[row.get("platform", "--"), row.get("error", "--")] for row in errors]
    _create_table(doc, ["平台", "原因"], rows)


def _add_data_source_section(doc: Document, report: Dict[str, Any]) -> None:
    doc.add_heading("六、数据来源声明", level=1)
    declaration = report.get("data_declaration") or "本报告所有数据均来自真实API调用，无模拟数据。"
    doc.add_paragraph(declaration)

    sources = report.get("data_sources", [])
    if sources:
        doc.add_paragraph("本次使用的数据源如下：")
        for source in sources:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(str(source))
    else:
        doc.add_paragraph("当前无已接通平台数据源。")

    generated_by = report.get("generated_by")
    report_time = report.get("report_time")
    footer = []
    if generated_by:
        footer.append(f"生成系统：{generated_by}")
    if report_time:
        footer.append(f"生成时间：{report_time}")
    if footer:
        doc.add_paragraph("；".join(footer))


def _collect_key_findings(report: Dict[str, Any]) -> List[str]:
    findings: List[str] = []
    statistics = report.get("overall_statistics", {})
    if statistics.get("critical_hosts", 0) > 0:
        findings.append(f"存在 {statistics.get('critical_hosts', 0)} 台严重状态主机，需要优先处理。")
    if statistics.get("warning_hosts", 0) > 0:
        findings.append(f"存在 {statistics.get('warning_hosts', 0)} 台警告状态主机，需要持续关注。")

    platforms = report.get("platforms", {})
    for platform in platforms.values():
        for item in platform.get("periodic_items", []):
            count = int(item.get("count", 0) or 0)
            if count > 0:
                findings.append(f"{platform.get('name', '--')} 的 {item.get('name', '--')} 发现 {count} 项问题。")

    return findings[:6]


def _build_periodic_recommendation(item: Dict[str, Any]) -> str:
    count = int(item.get("count", 0) or 0)
    if count == 0:
        return "建议动作：当前无需处置，保持现有巡检频率。"
    if item.get("code") == "snapshot":
        return "建议动作：优先清理超期快照，避免快照链持续增长影响性能与恢复窗口。"
    if item.get("code") == "naming":
        return "建议动作：补齐备注字段并统一命名规范，保证责任、系统和用途可追溯。"
    if item.get("code") == "idle_vm":
        return "建议动作：确认资产是否仍需保留，临时用途资产应按规则标识并及时回收。"
    if item.get("code") == "large_vm":
        return "建议动作：核查容量增长原因，评估归档、清理或扩容策略。"
    return "建议动作：结合巡检结果安排后续处置。"


def _create_table(doc: Document, headers: List[str], rows: Iterable[Iterable[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    return table


def _add_run(paragraph, text: str, size: float = 10.5, bold: bool = False):
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.font.bold = bold
    return run


def _format_percent(value: Any) -> str:
    try:
        return f"{float(value):.2f}%"
    except (TypeError, ValueError):
        return "0.00%"


def _format_report_date(report_date: str) -> str:
    if len(report_date) == 8 and report_date.isdigit():
        return f"{report_date[:4]}-{report_date[4:6]}-{report_date[6:]}"
    return report_date


def save_report(doc: Document, filepath: str):
    doc.save(filepath)


if __name__ == "__main__":
    document = create_inspection_report(company_name="厦门国际信托", report_date="20260413", data=None)
    save_report(document, "test_report.docx")
    print("测试报告已生成：test_report.docx")
